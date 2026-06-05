import os
from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

def create_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph()
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                p.add_run(line)

    doc.save(docx_path)

def sanitize_text(text):
    # Replace common unicode characters that cause issues with latin-1 fonts
    return text.replace('\u2013', '-').replace('\u2014', '--').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"').replace('\u2022', '*')

def create_pdf(md_path, pdf_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("helvetica", size=10)

    with open(md_path, 'r') as f:
        lines = f.readlines()

    for line in lines:
        line = sanitize_text(line.strip())
        if not line:
            pdf.ln(5)
            continue

        if line.startswith('# '):
            pdf.set_font("helvetica", "B", 16)
            pdf.cell(0, 10, line[2:], new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        elif line.startswith('## '):
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, line[3:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line.startswith('### '):
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 10, line[4:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif line.startswith('- '):
            pdf.set_font("helvetica", "", 10)
            pdf.multi_cell(190, 5, "* " + line[2:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font("helvetica", "", 10)
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        pdf.set_font("helvetica", "B", 10)
                    else:
                        pdf.set_font("helvetica", "", 10)
                    pdf.write(5, part)
                pdf.ln(5)
            else:
                pdf.multi_cell(190, 5, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(pdf_path)

if __name__ == "__main__":
    files_to_convert = [
        ('docs/applications/rws-ai-auditor/resume_alex_lomax.md', 'docs/applications/rws-ai-auditor/resume_alex_lomax'),
        ('docs/applications/rws-ai-auditor/resume_ahmed_ayyad.md', 'docs/applications/rws-ai-auditor/resume_ahmed_ayyad'),
        ('docs/applications/rws-ai-auditor/cover_letter_alex_lomax.md', 'docs/applications/rws-ai-auditor/cover_letter_alex_lomax'),
        ('docs/applications/rws-ai-auditor/cover_letter_ahmed_ayyad.md', 'docs/applications/rws-ai-auditor/cover_letter_ahmed_ayyad'),
    ]

    for md, base in files_to_convert:
        if os.path.exists(md):
            print(f"Converting {md}...")
            create_docx(md, base + ".docx")
            create_pdf(md, base + ".pdf")
        else:
            print(f"Skipping {md} - not found")
